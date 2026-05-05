"""
报告核心模块
提供报告生成器、报告类型、模板管理、导出功能和定时调度
"""

import logging
import uuid
import json
from typing import Optional, Dict, Any, List, Callable
from datetime import datetime, timedelta
from enum import Enum
from dataclasses import dataclass, field
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    
from jinja2 import Environment, FileSystemLoader, Template, TemplateNotFound

logger = logging.getLogger(__name__)


class ReportType(Enum):
    """报告类型枚举"""
    DAILY = "daily"           # 日巡检报告
    WEEKLY = "weekly"         # 周报
    MONTHLY = "monthly"       # 月报
    YEARLY = "yearly"         # 年报
    RCA = "rca"              # 根因分析报告
    CHANGE = "change"        # 变更报告
    SUMMARY = "summary"       # 总结报告
    CUSTOM = "custom"         # 自定义报告


class ExportFormat(Enum):
    """导出格式枚举"""
    HTML = "html"
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    MARKDOWN = "markdown"
    JSON = "json"


class ReportStatus(Enum):
    """报告状态枚举"""
    PENDING = "pending"       # 待生成
    GENERATING = "generating" # 生成中
    COMPLETED = "completed"   # 已完成
    FAILED = "failed"         # 生成失败
    ARCHIVED = "archived"     # 已归档


@dataclass
class ReportMetadata:
    """报告元数据"""
    report_id: str
    report_type: ReportType
    title: str
    description: str = ""
    author: str = "System"
    created_at: datetime = field(default_factory=datetime.now)
    period_start: Optional[datetime] = None
    period_end: Optional[datetime] = None
    tags: List[str] = field(default_factory=list)
    custom_fields: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Report:
    """报告对象"""
    metadata: ReportMetadata
    content: str = ""
    summary: str = ""
    attachments: List[str] = field(default_factory=list)
    status: ReportStatus = ReportStatus.PENDING
    error_message: Optional[str] = None
    generated_at: Optional[datetime] = None
    file_path: Optional[str] = None


class ReportScheduler:
    """报告调度器"""
    
    def __init__(self, scheduler):
        """
        初始化报告调度器
        
        Args:
            scheduler: TaskScheduler实例
        """
        self._scheduler = scheduler
        self._report_jobs: Dict[str, Dict[str, Any]] = {}
    
    def schedule_daily_report(
        self,
        report_generator: Callable,
        job_id: Optional[str] = None,
        hour: int = 8,
        minute: int = 0,
        **kwargs
    ) -> str:
        """调度日报生成"""
        job_id = job_id or f"daily_report_{uuid.uuid4().hex[:8]}"
        
        self._scheduler.add_cron_job(
            func=report_generator,
            cron=f"0 {minute} {hour} * * *",
            name=f"日报生成 - {job_id}",
            job_id=job_id,
            **kwargs
        )
        
        self._report_jobs[job_id] = {
            'type': ReportType.DAILY,
            'cron': f"0 {minute} {hour} * * *",
            'generator': report_generator,
        }
        
        logger.info(f"Scheduled daily report job: {job_id}")
        return job_id
    
    def schedule_weekly_report(
        self,
        report_generator: Callable,
        job_id: Optional[str] = None,
        day_of_week: int = 0,  # 0=周一
        hour: int = 9,
        minute: int = 0,
        **kwargs
    ) -> str:
        """调度周报生成"""
        job_id = job_id or f"weekly_report_{uuid.uuid4().hex[:8]}"
        
        self._scheduler.add_cron_job(
            func=report_generator,
            cron=f"0 {minute} {hour} * * {day_of_week}",
            name=f"周报生成 - {job_id}",
            job_id=job_id,
            **kwargs
        )
        
        self._report_jobs[job_id] = {
            'type': ReportType.WEEKLY,
            'cron': f"0 {minute} {hour} * * {day_of_week}",
            'generator': report_generator,
        }
        
        logger.info(f"Scheduled weekly report job: {job_id}")
        return job_id
    
    def schedule_monthly_report(
        self,
        report_generator: Callable,
        job_id: Optional[str] = None,
        day_of_month: int = 1,
        hour: int = 9,
        minute: int = 0,
        **kwargs
    ) -> str:
        """调度月报生成"""
        job_id = job_id or f"monthly_report_{uuid.uuid4().hex[:8]}"
        
        self._scheduler.add_cron_job(
            func=report_generator,
            cron=f"0 {minute} {hour} {day_of_month} * *",
            name=f"月报生成 - {job_id}",
            job_id=job_id,
            **kwargs
        )
        
        self._report_jobs[job_id] = {
            'type': ReportType.MONTHLY,
            'cron': f"0 {minute} {hour} {day_of_month} * *",
            'generator': report_generator,
        }
        
        logger.info(f"Scheduled monthly report job: {job_id}")
        return job_id
    
    def cancel_job(self, job_id: str):
        """取消调度任务"""
        if job_id in self._report_jobs:
            self._scheduler.remove_job(job_id)
            del self._report_jobs[job_id]
            logger.info(f"Cancelled report job: {job_id}")
    
    def get_scheduled_jobs(self) -> List[Dict[str, Any]]:
        """获取所有调度的报告任务"""
        return [
            {
                'job_id': job_id,
                'type': job['type'].value,
                'cron': job['cron'],
            }
            for job_id, job in self._report_jobs.items()
        ]


class TemplateRenderer:
    """模板渲染器"""
    
    def __init__(self, template_dirs: Optional[List[str]] = None):
        """
        初始化模板渲染器
        
        Args:
            template_dirs: 模板目录列表
        """
        self._template_dirs = template_dirs or ['templates']
        self._env = Environment(
            loader=FileSystemLoader(self._template_dirs),
            autoescape=True,
            trim_blocks=True,
            lstrip_blocks=True,
        )
        self._custom_filters: Dict[str, Callable] = {}
        self._register_default_filters()
    
    def _register_default_filters(self):
        """注册默认过滤器"""
        self._env.filters['datetime'] = lambda d, fmt='%Y-%m-%d %H:%M:%S': (
            d.strftime(fmt) if isinstance(d, datetime) else str(d)
        )
        self._env.filters['date'] = lambda d, fmt='%Y-%m-%d': (
            d.strftime(fmt) if isinstance(d, datetime) else str(d)
        )
        self._env.filters['duration'] = self._format_duration
        self._env.filters['percentage'] = lambda v, decimals=1: f"{v:.{decimals}f}%"
        self._env.filters['bytes'] = self._format_bytes
        self._env.filters['trend'] = self._format_trend
    
    @staticmethod
    def _format_duration(seconds: float) -> str:
        """格式化时长"""
        if seconds < 60:
            return f"{seconds:.0f}秒"
        elif seconds < 3600:
            return f"{seconds/60:.1f}分钟"
        elif seconds < 86400:
            return f"{seconds/3600:.1f}小时"
        else:
            return f"{seconds/86400:.1f}天"
    
    @staticmethod
    def _format_bytes(size: int) -> str:
        """格式化字节大小"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size < 1024:
                return f"{size:.1f}{unit}"
            size /= 1024
        return f"{size:.1f}PB"
    
    @staticmethod
    def _format_trend(value: float) -> str:
        """格式化趋势"""
        if value > 0:
            return f"↑ {value:.1f}%"
        elif value < 0:
            return f"↓ {abs(value):.1f}%"
        else:
            return "→ 0%"
    
    def add_filter(self, name: str, func: Callable):
        """添加自定义过滤器"""
        self._env.filters[name] = func
        self._custom_filters[name] = func
    
    def render(self, template_name: str, context: Dict[str, Any]) -> str:
        """
        渲染模板
        
        Args:
            template_name: 模板名称
            context: 渲染上下文
            
        Returns:
            渲染后的HTML字符串
        """
        try:
            template = self._env.get_template(template_name)
            return template.render(**context)
        except TemplateNotFound:
            logger.error(f"Template not found: {template_name}")
            raise
        except Exception as e:
            logger.error(f"Error rendering template {template_name}: {e}")
            raise
    
    def render_string(self, template_str: str, context: Dict[str, Any]) -> str:
        """渲染字符串模板"""
        try:
            template = Template(template_str)
            return template.render(**context)
        except Exception as e:
            logger.error(f"Error rendering template string: {e}")
            raise


class ReportExporter:
    """报告导出器"""
    
    def __init__(self, output_dir: str = "./reports"):
        """
        初始化导出器
        
        Args:
            output_dir: 输出目录
        """
        self._output_dir = Path(output_dir)
        self._output_dir.mkdir(parents=True, exist_ok=True)
        self._executor = ThreadPoolExecutor(max_workers=4)
    
    def export_html(self, report: Report, filename: Optional[str] = None) -> str:
        """导出为HTML"""
        filename = filename or f"{report.metadata.report_id}.html"
        filepath = self._output_dir / filename
        
        content = report.content
        if not content.startswith('<html'):
            content = self._wrap_html(content, report.metadata.title)
        
        filepath.write_text(content, encoding='utf-8')
        report.file_path = str(filepath)
        
        logger.info(f"Exported report to HTML: {filepath}")
        return str(filepath)
    
    def export_pdf(self, report: Report, filename: Optional[str] = None) -> str:
        """导出为PDF"""
        try:
            from weasyprint import HTML
        except ImportError:
            logger.warning("weasyprint not installed, falling back to HTML")
            return self.export_html(report, filename)
        
        filename = filename or f"{report.metadata.report_id}.pdf"
        filepath = self._output_dir / filename
        
        html_content = report.content
        if not html_content.startswith('<html'):
            html_content = self._wrap_html(html_content, report.metadata.title)
        
        HTML(string=html_content).write_pdf(str(filepath))
        report.file_path = str(filepath)
        
        logger.info(f"Exported report to PDF: {filepath}")
        return str(filepath)
    
    def export_word(self, report: Report, filename: Optional[str] = None) -> str:
        """导出为Word"""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, falling back to HTML")
            return self.export_html(report, filename)
        
        filename = filename or f"{report.metadata.report_id}.docx"
        filepath = self._output_dir / filename
        
        doc = Document()
        
        # 添加标题
        doc.add_heading(report.metadata.title, 0)
        
        # 添加内容（简单处理HTML）
        content = self._strip_html_tags(report.content)
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(str(filepath))
        report.file_path = str(filepath)
        
        logger.info(f"Exported report to Word: {filepath}")
        return str(filepath)
    
    def export_excel(self, report: Report, filename: Optional[str] = None) -> str:
        """导出为Excel"""
        try:
            import pandas as pd
            from openpyxl import Workbook
        except ImportError:
            logger.warning("pandas/openpyxl not installed, cannot export to Excel")
            return self.export_html(report, filename)
        
        filename = filename or f"{report.metadata.report_id}.xlsx"
        filepath = self._output_dir / filename
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Report Data"
        
        # 添加标题
        ws['A1'] = report.metadata.title
        
        # 解析JSON内容
        try:
            data = json.loads(report.content)
            if isinstance(data, list):
                for i, row in enumerate(data, start=3):
                    for j, val in enumerate(row.values() if isinstance(row, dict) else row, start=1):
                        ws.cell(row=i, column=j, value=str(val))
            elif isinstance(data, dict):
                for i, (key, val) in enumerate(data.items(), start=3):
                    ws.cell(row=i, column=1, value=key)
                    ws.cell(row=i, column=2, value=str(val))
        except (json.JSONDecodeError, TypeError):
            ws['A3'] = report.content
        
        wb.save(str(filepath))
        report.file_path = str(filepath)
        
        logger.info(f"Exported report to Excel: {filepath}")
        return str(filepath)
    
    def export_markdown(self, report: Report, filename: Optional[str] = None) -> str:
        """导出为Markdown"""
        filename = filename or f"{report.metadata.report_id}.md"
        filepath = self._output_dir / filename
        
        md_content = f"# {report.metadata.title}\n\n"
        md_content += f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        md_content += f"**报告类型**: {report.metadata.report_type.value}\n\n"
        
        if report.metadata.description:
            md_content += f"**描述**: {report.metadata.description}\n\n"
        
        md_content += "---\n\n"
        md_content += report.content
        
        filepath.write_text(md_content, encoding='utf-8')
        report.file_path = str(filepath)
        
        logger.info(f"Exported report to Markdown: {filepath}")
        return str(filepath)
    
    def export_json(self, report: Report, filename: Optional[str] = None) -> str:
        """导出为JSON"""
        filename = filename or f"{report.metadata.report_id}.json"
        filepath = self._output_dir / filename
        
        export_data = {
            'metadata': {
                'report_id': report.metadata.report_id,
                'report_type': report.metadata.report_type.value,
                'title': report.metadata.title,
                'description': report.metadata.description,
                'author': report.metadata.author,
                'created_at': report.metadata.created_at.isoformat(),
                'period_start': report.metadata.period_start.isoformat() if report.metadata.period_start else None,
                'period_end': report.metadata.period_end.isoformat() if report.metadata.period_end else None,
                'tags': report.metadata.tags,
                'custom_fields': report.metadata.custom_fields,
            },
            'content': report.content,
            'summary': report.summary,
            'status': report.status.value,
            'generated_at': report.generated_at.isoformat() if report.generated_at else None,
            'attachments': report.attachments,
        }
        
        filepath.write_text(json.dumps(export_data, ensure_ascii=False, indent=2), encoding='utf-8')
        report.file_path = str(filepath)
        
        logger.info(f"Exported report to JSON: {filepath}")
        return str(filepath)
    
    @staticmethod
    def _wrap_html(content: str, title: str) -> str:
        """包装HTML内容"""
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }}
        h2 {{ color: #555; margin-top: 30px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
        th {{ background-color: #007bff; color: white; }}
        tr:nth-child(even) {{ background-color: #f2f2f2; }}
        .alert {{ padding: 15px; background-color: #f8d7da; border: 1px solid #f5c6cb; border-radius: 4px; }}
        .success {{ background-color: #d4edda; border-color: #c3e6cb; }}
        .warning {{ background-color: #fff3cd; border-color: #ffeeba; }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
    
    @staticmethod
    def _strip_html_tags(html: str) -> str:
        """去除HTML标签"""
        import re
        clean = re.sub('<[^>]+>', '', html)
        return clean.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>')


class ReportGenerator:
    """
    报告生成器核心类
    
    功能：
    - 多种类型报告生成
    - 模板渲染
    - 多格式导出
    - 定时调度
    """
    
    def __init__(
        self,
        config: Optional[Dict[str, Any]] = None,
        template_dirs: Optional[List[str]] = None,
        output_dir: str = "./reports"
    ):
        """
        初始化报告生成器
        
        Args:
            config: 配置字典
            template_dirs: 模板目录列表
            output_dir: 输出目录
        """
        self._config = config or {}
        self._template_renderer = TemplateRenderer(template_dirs)
        self._exporter = ReportExporter(output_dir)
        self._reports_store: Dict[str, Report] = {}
        self._scheduler: Optional[ReportScheduler] = None
    
    def set_scheduler(self, scheduler):
        """设置调度器"""
        self._scheduler = ReportScheduler(scheduler)
    
    def create_report(
        self,
        report_type: ReportType,
        title: str,
        description: str = "",
        author: str = "System",
        period_start: Optional[datetime] = None,
        period_end: Optional[datetime] = None,
        tags: Optional[List[str]] = None,
        custom_fields: Optional[Dict[str, Any]] = None,
    ) -> Report:
        """
        创建报告对象
        
        Args:
            report_type: 报告类型
            title: 标题
            description: 描述
            author: 作者
            period_start: 统计周期开始
            period_end: 统计周期结束
            tags: 标签
            custom_fields: 自定义字段
            
        Returns:
            Report对象
        """
        report_id = f"{report_type.value}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
        
        metadata = ReportMetadata(
            report_id=report_id,
            report_type=report_type,
            title=title,
            description=description,
            author=author,
            period_start=period_start or datetime.now() - timedelta(days=1),
            period_end=period_end or datetime.now(),
            tags=tags or [],
            custom_fields=custom_fields or {},
        )
        
        report = Report(metadata=metadata)
        self._reports_store[report_id] = report
        
        return report
    
    def generate(
        self,
        report: Report,
        template_name: Optional[str] = None,
        context: Optional[Dict[str, Any]] = None,
    ) -> Report:
        """
        生成报告
        
        Args:
            report: 报告对象
            template_name: 模板名称
            context: 渲染上下文
            
        Returns:
            生成的报告
        """
        report.status = ReportStatus.GENERATING
        
        try:
            context = context or {}
            context['metadata'] = report.metadata
            
            if template_name:
                report.content = self._template_renderer.render(template_name, context)
            elif context.get('content'):
                report.content = context['content']
            else:
                report.content = self._generate_default_content(report)
            
            report.summary = self._generate_summary(report)
            report.status = ReportStatus.COMPLETED
            report.generated_at = datetime.now()
            
            logger.info(f"Generated report: {report.metadata.report_id}")
            
        except Exception as e:
            report.status = ReportStatus.FAILED
            report.error_message = str(e)
            logger.error(f"Failed to generate report {report.metadata.report_id}: {e}")
        
        return report
    
    def export(
        self,
        report: Report,
        format: ExportFormat,
        filename: Optional[str] = None,
    ) -> str:
        """
        导出报告
        
        Args:
            report: 报告对象
            format: 导出格式
            filename: 文件名
            
        Returns:
            导出文件路径
        """
        if report.status != ReportStatus.COMPLETED:
            raise ValueError(f"Report {report.metadata.report_id} is not completed")
        
        export_funcs = {
            ExportFormat.HTML: self._exporter.export_html,
            ExportFormat.PDF: self._exporter.export_pdf,
            ExportFormat.WORD: self._exporter.export_word,
            ExportFormat.EXCEL: self._exporter.export_excel,
            ExportFormat.MARKDOWN: self._exporter.export_markdown,
            ExportFormat.JSON: self._exporter.export_json,
        }
        
        if format not in export_funcs:
            raise ValueError(f"Unsupported export format: {format}")
        
        return export_funcs[format](report, filename)
    
    def _generate_default_content(self, report: Report) -> str:
        """生成默认报告内容"""
        meta = report.metadata
        
        content = f"<h1>{meta.title}</h1>\n"
        content += f"<p><strong>报告ID:</strong> {meta.report_id}</p>\n"
        content += f"<p><strong>类型:</strong> {meta.report_type.value}</p>\n"
        content += f"<p><strong>生成时间:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>\n"
        
        if meta.description:
            content += f"<p><strong>描述:</strong> {meta.description}</p>\n"
        
        if meta.period_start and meta.period_end:
            content += f"<p><strong>统计周期:</strong> {meta.period_start.strftime('%Y-%m-%d')} 至 {meta.period_end.strftime('%Y-%m-%d')}</p>\n"
        
        return content
    
    def _generate_summary(self, report: Report) -> str:
        """生成报告摘要"""
        meta = report.metadata
        
        summary = f"本报告为{meta.report_type.value}类型报告"
        
        if meta.period_start and meta.period_end:
            duration = (meta.period_end - meta.period_start).total_seconds()
            if duration < 86400:
                summary += f"，统计周期为{duration/3600:.1f}小时"
            else:
                summary += f"，统计周期为{duration/86400:.1f}天"
        
        return summary
    
    def get_report(self, report_id: str) -> Optional[Report]:
        """获取报告"""
        return self._reports_store.get(report_id)
    
    def list_reports(
        self,
        report_type: Optional[ReportType] = None,
        status: Optional[ReportStatus] = None,
        limit: int = 100,
    ) -> List[Report]:
        """列出报告"""
        reports = list(self._reports_store.values())
        
        if report_type:
            reports = [r for r in reports if r.metadata.report_type == report_type]
        
        if status:
            reports = [r for r in reports if r.status == status]
        
        reports.sort(key=lambda r: r.metadata.created_at, reverse=True)
        
        return reports[:limit]
    
    def delete_report(self, report_id: str) -> bool:
        """删除报告"""
        if report_id in self._reports_store:
            del self._reports_store[report_id]
            return True
        return False
    
    @property
    def scheduler(self) -> Optional[ReportScheduler]:
        """获取调度器"""
        return self._scheduler
